"""DOCX parser using python-docx."""

from typing import Any

from docx import Document

from policymind.documents.models import DocumentBlock, ParsedDocument


class DOCXParser:
    """Extract text, headings, and tables from .docx files."""

    supported_mime_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    def parse(self, content: bytes, filename: str = "") -> ParsedDocument:
        from io import BytesIO

        doc = Document(BytesIO(content))
        blocks: list[DocumentBlock] = []
        page_count = 1  # python-docx does not expose page count directly

        heading_stack: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break

                if para is None:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    heading_stack = _update_heading_stack(heading_stack, style_name, text)
                    blocks.append(
                        DocumentBlock(
                            text=text,
                            page_number=page_count,
                            block_type="title",
                            heading_path=tuple(heading_stack),
                        )
                    )
                else:
                    blocks.append(
                        DocumentBlock(
                            text=text,
                            page_number=page_count,
                            block_type="text",
                            heading_path=tuple(heading_stack),
                        )
                    )

            elif tag == "tbl":
                # Extract table as structured text
                table_text = _extract_table_text(element)
                if table_text:
                    blocks.append(
                        DocumentBlock(
                            text=table_text,
                            page_number=page_count,
                            block_type="table",
                            heading_path=tuple(heading_stack),
                        )
                    )

        return ParsedDocument(
            title=filename or "Untitled",
            page_count=page_count,
            blocks=blocks,
            metadata={"parser": "python-docx"},
        )


def _update_heading_stack(
    stack: list[str], style_name: str, title: str
) -> list[str]:
    level = 1
    for char in style_name:
        if char.isdigit():
            level = int(char)
            break
    stack = stack[: level - 1]
    stack.append(title)
    return stack


def _extract_table_text(element: Any) -> str:
    """Extract table rows as TSV-like text."""
    rows: list[str] = []
    for row in element.iterchildren():
        tag = row.tag.split("}")[-1] if "}" in row.tag else row.tag
        if tag != "tr":
            continue
        cells: list[str] = []
        for cell in row.iterchildren():
            cell_text = "".join(
                node.text or ""
                for node in cell.iterdescendants()
                if node.text
            )
            cells.append(cell_text.strip())
        rows.append(" | ".join(cells))
    return "\n".join(rows)
